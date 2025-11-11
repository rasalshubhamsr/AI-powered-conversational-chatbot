#!/usr/bin/env python3
"""
Generate or overwrite a Word document (doc.docx) from a plain text file.

Usage:
  python scripts/generate_docx.py [input_text_file]

- If input_text_file is not provided, defaults to scripts/content_beginner_guide.txt
- Output is written to ./doc.docx (overwritten if exists)
"""
import os
import sys
from docx import Document
from docx.shared import Pt

DEFAULT_INPUT = os.path.join("scripts", "content_beginner_guide.txt")
OUTPUT_DOCX = os.path.join("doc.docx")

def write_docx_from_text(input_path: str, output_path: str) -> None:
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"Input text file not found: {input_path}")

    with open(input_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Set a base style (optional)
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Add a title
    doc.add_heading("Beginner Guide: How We Built This Project", level=1)

    # Add paragraphs from content, preserving empty lines
    for line in content.splitlines():
        if line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    doc.save(output_path)
    print(f"Wrote: {output_path}")


def main():
    input_path = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_INPUT
    write_docx_from_text(input_path, OUTPUT_DOCX)

if __name__ == "__main__":
    main()
